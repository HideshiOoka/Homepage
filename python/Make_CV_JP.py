#%%
import pandas as pd

from docx import Document
LANG = 0 # English:0 or Japanese:1
# from docx.shared import Pt

def sort(authors):
    sorted_authors = ""
    for a in authors.split(" and "):
        last,first  = a.split(", ")
        sorted_authors += first + " " + last + ", "
    sorted_authors = sorted_authors[:-2]        
    return sorted_authors
def add_formatted_title(p,title):  
    title = title.replace("--","-")   # -- in SPET
    if "CO" not in p.text and "MoS2" not in p.text:
        p.add_run(title)
        return p
    else:
        # break the runs? or do something?# # CO2, MoS2        
        p.add_run(title) # need some kind of formatting
        return p

def make_publication_list(bib):
    doc.add_heading("【論文リスト】", level=1) # reorder if original/reviews need to be separated?
    with open(bib, encoding = "UTF-8") as f:
        bib_data = f.read().replace("\n", "") # leave doi and other information in the bib data. comments might be beneficial for biorxiv and books.
    entries = bib_data.split("@")[1:]
    num_publications =len(entries)
    corresponding_authors = str(bib_data.count("Ooka*"))
    first_authors = str(bib_data.count("@article{Ooka"))
    for i, entry in enumerate(entries):
        authors = entry.split("author = {")[1].split("},")[0]
        authors = sort(authors)
        journal = entry.split("journal = {")[1].split("},   ")[0]
        year = entry.split("year = {")[1].split("},   ")[0]
        volume = entry.split("volume = {")[1].split("},   ")[0]
        pages = entry.split("pages = {")[1].split("},   ")[0].replace("--","-")
        title = entry.split("title = {")[1].split("}}")[0]
        p = doc.add_paragraph(f"{i+1}.  ")
        p.add_run(authors)
        p.add_run(" \"")
        p = add_formatted_title(p,title) 
        p.add_run("\", ")
        p.add_run(journal).italic = True
        p.add_run(", ")
        p.add_run(year).bold = True
        p.add_run(", ")
        p.add_run(pages)
        p.add_run(".").add_break()
def make_funding_list(funding_csv):
    doc.add_heading("【外部資金獲得歴】", level=1)
    df = pd.read_csv(funding_csv)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i]
        start,finish,title, name, source, PI, amount, unit = data[["Start","Finish","Title_JP","PJ_Name_JP","Funding_Source_JP","PI","Amount","Unit"]]
        p = doc.add_paragraph("")
        if PI == "PI":
            PI = "(研究代表者)"
        else:
            PI = "(研究分担者)"    
        p.add_run(f"{i+1}.  {source} {name} {PI}\n").bold = True
        p.add_run(f"課題名：{title}\n")
        p.add_run(f"研究期間：{start} - {finish}\n")
        p.add_run(f"研究費総額：{amount}円").add_break()
        # funding amount with comma?

def make_award_list(award_csv):
    doc.add_heading("【受賞歴】", level=1)
    df = pd.read_csv(award_csv)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i]
        date,prize,origin = data[["Date","Prize_JP","From_JP"]]
        p = doc.add_paragraph("")
        p.add_run(f"{date},{prize},{origin}").add_break()

def make_patent_list(patent_csv):
    doc.add_heading("【知財・特許】", level=1)
    df = pd.read_csv(patent_csv)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i]
        date,status,authors,title,ID = data
        p = doc.add_paragraph("")
        p.add_run(f"{i+1}.  {title},{authors},{status}").add_break()

def format_date(date,n=8): # change format if n < 8
    formatted_date = date[:4]+"."+date[4:6]+date[6:8]
    formatted_date = formatted_date.replace("X","")
    return formatted_date

def make_press_list(press_csv):
    pass ##########
def make_presentation_list(presentation_csv):
    doc.add_heading("【発表リスト】", level=1)
    df = pd.read_csv(presentation_csv)
    N = df.shape[0]
    ### sort df according to invited/oral/poster
    ### international/domestic
    for i in range(N):
        data = df.iloc[i]   
        date,conference,venue,location,title,authors,format,notes = data[:8]
        
        formatted_date = format_date(date)
        

        p = doc.add_paragraph("")
        p.add_run(f"{i+1}.  {authors} ")
        p.add_run(f'"{title}"').bold = True
        p.add_run(f"{conference}, {venue} ({formatted_date})").add_break()

items = ["Heading", "Publications", "Presentations","Patents","Press","Funding","Awards"]
doc = Document("CV_template.docx")
    

for item in items:
    doc.add_paragraph(item,style = "Heading 1")
make_publication_list("Publications.bib")
make_funding_list("Funding.csv")
make_award_list("Awards.csv")
make_patent_list("Patents.csv")
# make_presentation_list("Presentations.csv")
# make_press_list("Press.csv")

doc.save("CV.docx")
