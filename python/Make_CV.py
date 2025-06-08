#%%
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import RGBColor, Pt, Cm
from docx.enum.text import WD_LINE_SPACING
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from CV_utils import sort_authors #,quote
from update_publications import reorder_publications
# date 20171101 for the CSRS Interim report might be off
#%%
def set_col_widths(table, widths = [Cm(1.5), Cm(14.5)]):
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width            

def h1(txt):
    yellow_background = OxmlElement("w:shd")
    yellow_background.set(qn("w:fill"), "#FFF2CC")
    tbl = doc.add_table(1,1)
    tbl.cell(0,0).text = txt
    tbl.rows[0].cells[0]._tc.get_or_add_tcPr().append(yellow_background)
    header = tbl.cell(0,0).paragraphs[0].runs[0]
    header.font.size = Pt(14)
    bold(header)
    """
    header = tbl.cell(0,0).paragraphs[0].runs[0].font.name = "Arial"
    if LANG == "_JP":
        header = tbl.cell(0,0).paragraphs[0].runs[0].font.name = "游ゴシック" # Doesn't work, not sure why
    header = tbl.cell(0,0).paragraphs[0].runs[0].font.bold = True
    """
    # tbl.rows.height = Pt(64)
    # tbl.line_spacing = WD_LINE_SPACING.MULTIPLE(5) # Pt(18) # WD_LINE_SPACING.EXACTLY

def bold(txt, font_name = "Arial"):
    txt.font.bold = True
    txt.font.name = font_name
    if txt.text.isascii() == False:
        font_name = "游ゴシック"
        txt.font.name = font_name
        txt._element.rPr.rFonts.set(qn('w:eastAsia'), txt.font.name)


def red(txt, color = "B10026"):
    txt.font.color.rgb = RGBColor.from_string(color)

def h2(text):
    p = doc.add_paragraph()
    header = p.add_run(text)
    header.font.size = Pt(14)
    header.font.underline = True
    bold(header)
    
def add_formatted_title(p,title):
    # 20240820 stopped making title bold. Instead, use bold on year and journal
    quotes = ['"','"']
    if title.isascii() == False:
        quotes = ['「','」']
    quote = p.add_run(f" {quotes[0]}")
    # bold(quote)
    title = title.replace("--","-")   # -- in SPET
    if "$_" not in title: # no subscripts
        title = p.add_run(title)
        # bold(title)
    else: 
        substrings = title.split(r"$_")
        for substring in substrings:
            if "$" in substring:
                subscript, upright = substring.split("$")
            else:
                subscript, upright = "", substring
            sub_text = p.add_run(subscript)
            sub_text.font.subscript = True
            # bold(sub_text)
            upright = p.add_run(upright)
            # bold(upright)
    quote = p.add_run(f"{quotes[1]}")
    # bold(quote)    
    p.add_run(" ") # assuming that there is some information coming after this 
    return p

def add_formatted_authors(p,authors):
    my_name = "Hideshi Ooka"
    try:
        before, after = authors.split(my_name)
    except ValueError:
        if LANG == "_JP":
            my_name = "大岡英史"
            before, after = authors.split(my_name)
        else:
            print(authors)
            before, after = "XXX", "XXX"
    text_tag_list = [f"{before};",
                     f"{my_name};bu",
                     f"{after};"]
    # p = write_docx(text_tag_list, add_paragraph = False)
    p.add_run(before)
    me = p.add_run(my_name)
    me.font.underline = True
    bold(me)
    p.add_run(after)
    return p

def write_entry(i, entry): # need to write the df     
    tbl = doc.add_table(1,2)
    tbl.allow_autofit = False
    tbl.autofit = False
    tbl.cell(0,0).text = f"{i+1}."
    tbl.cell(0,1).text = ""
    p = tbl.cell(0,1).paragraphs[0]
    
    type,notes,status,url,preprint,bib_key,j_key,title,pages,volume,year,fullname,abbrv,journal,authors,ENTRYTYPE,ID,priority = entry
    # print(authors)
    # ID, authors, journal, year, volume, pages, doi, notes, ENTRYTYPE, type, title, fullname, abbrv, status = entry
    year = str(int(year))
    if volume !="":
        volume = str(int(volume))
    authors = sort_authors(authors)
    p = add_formatted_authors(p,authors) ### Authors:
    p = add_formatted_title(p,title) 
    journal = p.add_run(journal)
    journal.italic = True
    bold(journal)
    p.add_run(", ")
    year = p.add_run(year)
    bold(year)
    p.add_run(", ")
    if pages != "":
        pages = pages.replace("--","-")
        p.add_run(volume).italic = True
        p.add_run(", ")
        p.add_run(pages)
    else: # it doesn't have pages
        # p.add_run(doi) # it must have doi
        p.add_run("(")
        p.add_run(status).italic = True
        p.add_run(")")
    p.add_run(".").add_break()
    if notes != "":
        if LANG == "_JP":
            notes = notes.replace("Representative Paper", "代表論文")
        comment = p.add_run(notes)
        bold(comment)
        # comment.font.underline = True
        comment.font.color.rgb = RGBColor.from_string("B10026")
        comment.add_break()
    set_col_widths(tbl)
    
def make_publication_list(csv_file, separate_reviews = False):
    publications = pd.read_csv(csv_file, index_col = 0, encoding_errors = "backslashreplace")
    publications = reorder_publications(publications)

    # The 2 step sort is done to prioritize empty status (=accepted) articles in the order
    publications = publications.sort_values(by = ["type","year"], ascending = [True, False]).fillna("")
    original = publications[publications["type"]=="Original"]
    review_others = publications[publications["type"]!="Original"]

    num_corresponding = publications["author"].str.contains("Ooka\*").sum() # Must write Ooka*+, not Ooka+*
    num_first = publications["ID"].str.contains("Ooka").sum() + publications["author"].str.contains("Ooka+",regex=False).sum() + publications["author"].str.contains("Ooka*+",regex=False).sum()
    out_html = f"<br>(Corresponding Author: {num_corresponding}, First Author: {num_first}, +: Dual First Author)<br>\n\n"

    df_list =[original, review_others]
    header = "Academic Publications (All Peer Reviewed)"
    sub_headers = ["Original Papers","Reviews"]
    unit = ""
    #######################
    if LANG == "_JP":
        header = "学術論文 (査読あり)"
        sub_headers = ["原著論文","総説"]
        unit = " 報"
    h1(header) # doc.add_heading(header, level=1)
    for i, df in enumerate(df_list):
        N = df.shape[0]
        h2(f"{sub_headers[i]}: {N}{unit}")
        for j in range(N):
            entry = df.iloc[j]
            write_entry(j, entry)
    p = doc.add_paragraph("\n")

def format_date(date,n=8): # change format if n < 8
    date = str(date)
    formatted_date = date[:4]+"/"+date[4:6]+"/"+date[6:8]
    formatted_date = formatted_date.replace("X","")
    return formatted_date

months = ["","January","February","March","April","May","June","July","August","September","October","November","December"]
def make_funding_list(funding_csv):
    header = "Funding (Japanese Titles were Translated to English)"
    PI_text = "Principal Investigator"
    Co_PI_text = "Co-Investigator"
    sub_headers = [PI_text, Co_PI_text]
    unit = "yen"
    if LANG == "_JP":
        header = "外部資金獲得状況"
        PI_text = "研究代表者"
        Co_PI_text = "研究分担者"
        unit = "円"
        sub_headers = [f"【{PI_text}】", f"【{Co_PI_text}】"]
    h1(header) # doc.add_heading(header, level=1)
    all_df = pd.read_csv(funding_csv, encoding_errors = "backslashreplace")
    all_df = all_df.sort_values(by = ["Start"], ascending = False)
    PI_df = all_df[all_df["PI"] == "PI"]
    Co_PI_df = all_df[all_df["PI"] != "PI"]

    dfs = [PI_df, Co_PI_df]
    
    for j, df in enumerate(dfs):
        # tbl = doc.add_table(1,1)
        # tbl.cell(0,0).text = sub_headers[j]
        h2(sub_headers[j])

        N = df.shape[0]
        for i in range(N):
            data = df.iloc[i]
            start,finish,title, name, source, PI, amount, unit_ = data[["Start","Finish","Title","PJ_Name","Funding_Source","PI","Amount","Unit"]] # change variable name later
            if PI == "PI":
                PI = PI_text
            else:
                PI = Co_PI_text
                
            tbl = doc.add_table(1,2)
            tbl.allow_autofit = False
            tbl.autofit = False
            tbl.cell(0,0).text = f"{i+1}."
            tbl.cell(0,1).text = ""
            p = tbl.cell(0,1).paragraphs[0]
            # p.add_run(f"{source} {name} ({PI})\n")
            p.add_run(f"{source} {name}\n")
            p = add_formatted_title(p,title) 
            start = str(start)
            finish = str(finish)
            formatted_start = start[:4]# + " " + months[int(start[4:])]
            formatted_finish = finish[:4]# + " " +months[int(finish[4:])]
            p.add_run(f"\n({formatted_start} - {formatted_finish}, {amount} {unit})").add_break()
            set_col_widths(tbl) 
    p = doc.add_paragraph("\n")

def make_award_list(award_csv):
    header = "Awards (Japanese Titles were Translated to English)"
    if LANG == "_JP":
        header = "受賞歴"
    h1(header) # doc.add_heading(header, level=1)
    df = pd.read_csv(award_csv, encoding_errors = "backslashreplace")
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i].astype(str)
        date,prize,origin,notes = data[["Date","Prize","From","Notes"]]
        date = str(date)
        date = f"{date[:4]}/{date[4:6]}/{date[6:]}"

        tbl = doc.add_table(1,2)
        tbl.allow_autofit = False
        tbl.autofit = False
        tbl.cell(0,0).text = f"{i+1}."
        tbl.cell(0,1).text = ""
        p = tbl.cell(0,1).paragraphs[0]

        prize = p.add_run(f"{prize}")
        bold(prize)
        p.add_run(f", {origin} ({date}).").add_break()
        if notes !="nan":
            comment = p.add_run(notes)
            bold(comment)
            # comment.font.underline = True
            comment.font.color.rgb = RGBColor.from_string("B10026")
            comment.add_break()
        set_col_widths(tbl)            
    p = doc.add_paragraph("\n")

def make_patent_list(patent_csv):
    header = "Patents"
    if LANG == "_JP":
        header = "知財・特許"
    h1(header) # doc.add_heading(header, level=1)
    df = pd.read_csv(patent_csv, encoding_errors = "backslashreplace")
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i,:5]
        date,status,authors,title,ID = data
        authors = authors.replace(" and",",")
        tbl = doc.add_table(1,2)
        tbl.allow_autofit = False
        tbl.autofit = False
        tbl.cell(0,0).text = f"{i+1}."
        tbl.cell(0,1).text = ""
        p = tbl.cell(0,1).paragraphs[0]
    
        ### Authors:
        p = add_formatted_authors(p,authors)
        p = add_formatted_title(p,title) 
        p.add_run(f" {ID} ({status}).").add_break()
        set_col_widths(tbl) 
    p = doc.add_paragraph("\n")


def make_presentation_list(presentation_csv):
    header = "Presentations (Japanese Titles were Translated to English)"
    if LANG == "_JP":
        header = "学会発表"
    h1(header) # doc.add_heading(header, level=1)
    format_list = ["Invited","Oral","Poster"]
    format_list_JP = ["【招待講演】","【口頭発表】","【ポスター発表】"]
    all_df = pd.read_csv(presentation_csv, encoding_errors = "backslashreplace")
    for j,df in enumerate(format_list):
        df = all_df[all_df["Format"] == format_list[j]]
        df = df.sort_values(by = "Date", ascending = False)
        N = df.shape[0]
        sub_header = format_list[j] + " Presentations"+f" ({N})"
        if LANG == "_JP":
            sub_header = format_list_JP[j]
        h2(sub_header) # doc.add_heading(sub_header, level=2)
        
        for i in range(N):
            data = df.iloc[i].astype(str)
            date,conference,venue,country_or_city,title,authors,format,notes = data[:8]
            formatted_date = format_date(date)
            tbl = doc.add_table(1,2)
            tbl.allow_autofit = False
            tbl.autofit = False
            tbl.cell(0,0).text = f"{i+1}."
            tbl.cell(0,1).text = ""
            p = tbl.cell(0,1).paragraphs[0]

            p = add_formatted_authors(p,authors)
            p = add_formatted_title(p,title) 
            p.add_run(f"{conference}, {venue}, {country_or_city} ({formatted_date}).").add_break()
            if notes !="nan":
                comment = p.add_run(notes)
                print(notes)
                bold(comment)
                # comment.font.underline = True
                red(comment)
                comment.add_break()
            set_col_widths(tbl)
    p = doc.add_paragraph("\n")  

def make_press_list(press_csv):
    pass ##########

def make_education_list(education_csv):
    pass

def make_others_list(others_csv):
    # TOEIC on 20240421
    # CSRS diversity
    # Finished Globis leadership
    pass
def make_others_list(others_csv):
    header = "Others"
    if LANG == "_JP":
        header = "その他"
    h1(header) # doc.add_heading(header, level=1)
    df = pd.read_csv(others_csv)
    df = df.sort_values(by = "Date", ascending = False)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i].astype(str)
        date,contents = data
        formatted_date = format_date(date)
        p = doc.add_paragraph(f"{i+1}.  ")
        p.add_run(f"{contents}.").add_break()
for LANG in ["","_JP"]:
        
    doc = Document(f"../templates/CV_Template{LANG}.docx")
    make_publication_list(f"../achievements/Publications.csv", separate_reviews=True)
    make_presentation_list(f"../achievements/Presentations{LANG}.csv")
    make_patent_list(f"../achievements/Patents{LANG}.csv")
    make_award_list(f"../achievements/Awards{LANG}.csv")
    make_funding_list(f"../achievements/Funding{LANG}.csv")
    make_education_list(f"../achievements/Funding{LANG}.csv")
    # make_press_list("../achievements/Press.csv")
    # make_others_list("../achievements/Others{LANG}.csv")
    doc.save(f"../achievements/Ooka_CV{LANG}_draft.docx")




#%% To do list:
# change order in the exported publications
# The draft version has equal width columns and needs to be changed manually. 20240820
# Also check new line locations
# After outputting the docs, the 1ページの行数を指定時に文字をグリッド線に合わせる needs to be unchecked (Ctrl + A). 20240820, not necessary?

# 年度
# Needs Japanese mode in general

def write_docx(text_tag_list, add_paragraph = True):
    if add_paragraph == True:
        p = doc.add_paragraph()
    for text_tag in text_tag_list:
        text, tags = text_tag.split(";")
        new_text = p.add_run(text)
        if "b" in tags:
            new_text.font.bold = True
        if "i" in tags:
            new_text.font.italic = True
        if "u" in tags:
            new_text.font.underline = True
        if "^" in tags:
            new_text.font.superscript = True
        if "_" in tags:
            new_text.font.subscript = True            
    return p

text_tag_list = ["abc;","abc;bu_","abssc;iu^"]
doc = Document(f"../templates/CV_Template.docx")
write_docx(text_tag_list)
doc.save(f"../achievements/test.docx")

