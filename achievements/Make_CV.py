#%%
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import RGBColor
LANG = 0 # English:0 or Japanese:1
# date 20171101 for the CSRS Interim report might be off
# from docx.shared import Pt
# search Km and $ and { for TeX

def sort(authors):
    sorted_authors = ""
    for a in authors.split(" and "):
        last,first  = a.split(", ")
        sorted_authors += first + " " + last + ", "
    sorted_authors = sorted_authors[:-2]        
    return sorted_authors
def add_formatted_title(p,title):  
    title = title.replace("--","-")   # -- in SPET
    if "CO" not in p.text and "MoS2" not in p.text:
        p.add_run(title).bold = True
        return p
    else:
        # break the runs? or do something?# # CO2, MoS2        
        p.add_run(title).bold = True # need some kind of formatting
        return p

def add_formatted_authors(p,authors):
    before, after = authors.split("Hideshi Ooka")
    p.add_run(before)
    me = p.add_run("Hideshi Ooka")
    me.font.underline = True
    me.font.bold = True
    p.add_run(after)
    return p
def make_publication_list(bib):
    doc.add_heading("Academic Publications (All Peer Reviewed)", level=1) # reorder if original/reviews need to be separated?
    with open(bib, encoding = "UTF-8") as f:
        bib_data = f.read().replace("\n", "") # leave doi and other information in the bib data. comments might be beneficial for biorxiv and books.
    bib_data = bib_data.replace("XXX","")        
    entries = bib_data.split("@")[1:]
    num_publications =len(entries)
    corresponding_authors = str(bib_data.count("Ooka*"))
    first_authors = str(bib_data.count("@article{Ooka"))
    for i, entry in enumerate(entries):
        entry = entry.split("}}")[0]
        authors = entry.split("author = {")[1].split("},")[0]
        authors = sort(authors)
        journal = entry.split("journal = {")[1].split("},   ")[0]
        year = entry.split("year = {")[1].split("},   ")[0]
        volume = entry.split("volume = {")[1].split("},   ")[0]
        pages = entry.split("pages = {")[1].split("},   ")[0].replace("--","-")
        title = entry.split("title = {")[1].split("},   ")[0]
        notes,doi,status ="","",""
        if "notes" in entry:
            notes = entry.split("notes = {")[1].split("},   ")[0]
        if "doi" in entry:
            doi = entry.split("doi = {")[1].split("},")[0]
        if "status" in entry:            
            status = entry.split("status = {")[1].split("},")[0]
        p = doc.add_paragraph(f"{i+1}.  ")
        ### Authors:
        p = add_formatted_authors(p,authors)
        p.add_run(" \"").bold = True
        p = add_formatted_title(p,title) 
        p.add_run("\" ").bold = True
        p.add_run(journal).italic = True
        p.add_run(", ")
        p.add_run(year).bold = True
        p.add_run(", ")
        if volume != "":
            p.add_run(volume).italic = True
            p.add_run(", ")
        if pages != "":
            p.add_run(pages)
        else: # it doesn't have pages
            p.add_run(doi) # it must have doi
            p.add_run(" (")
            p.add_run(status).italic = True
            p.add_run(")")
        p.add_run(".").add_break()
        if notes != "":
            comment = p.add_run(notes)
            comment.font.bold = True
            comment.font.name = "Arial"
            # comment.font.underline = True
            comment.font.color.rgb = RGBColor.from_string("B10026")
            comment.add_break()


months = ["","January","February","March","April","May","June","July","August","September","October","November","December"]
def make_funding_list(funding_csv):
    doc.add_heading("Funding (Japanese Titles were Translated to English)", level=1)
    df = pd.read_csv(funding_csv)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i]
        start,finish,title, name, source, PI, amount, unit = data[["Start","Finish","Title","PJ_Name","Funding_Source","PI","Amount","Unit"]]
        if PI == "PI":
            PI = "Principal Investigator"
        else:
            PI = "Co-Investigator"    
        p = doc.add_paragraph(f"{i+1}.  ")
        p.add_run(f"{source} {name} ({PI})\n")
        p.add_run(" \"").bold = True
        p = add_formatted_title(p,title) 
        p.add_run("\" ").bold = True
        start = str(start)
        finish = str(finish)
        formatted_start = start[:4] + " " + months[int(start[4:])]
        formatted_finish = finish[:4] + " " +months[int(finish[4:])]
        p.add_run(f"({formatted_start} - {formatted_finish}, {amount} yen)").add_break()

def make_award_list(award_csv):
    doc.add_heading("Awards (Japanese Titles were Translated to English)", level=1)
    df = pd.read_csv(award_csv)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i].astype(str)
        date,prize,origin,notes = data[["Date","Prize","From","Notes"]]
        date = str(date)
        date = f"{date[:4]}/{date[4:6]}/{date[6:]}"
        p = doc.add_paragraph(f"{i+1}.  ")
        p.add_run(f"{prize}").bold = True
        p.add_run(f", {origin} ({date}).").add_break()
        if notes !="nan":
            print(notes)
            comment = p.add_run(notes)
            comment.font.bold = True
            # comment.font.underline = True
            comment.font.color.rgb = RGBColor.from_string("B10026")
            comment.add_break()

def make_patent_list(patent_csv):
    doc.add_heading("Patents", level=1)
    df = pd.read_csv(patent_csv)
    N = df.shape[0]
    for i in range(N):
        data = df.iloc[i,:5]
        date,status,authors,title,ID = data
        authors = authors.replace(" and",",")
        p = doc.add_paragraph(f"{i+1}.  ")
        ### Authors:
        p = add_formatted_authors(p,authors)
        p.add_run(" \"")
        p = add_formatted_title(p,title) 
        p.add_run("\", ")
        p.add_run(f"{ID} ({status}).").add_break()

def format_date(date,n=8): # change format if n < 8
    date = str(date)
    formatted_date = date[:4]+"/"+date[4:6]+"/"+date[6:8]
    formatted_date = formatted_date.replace("X","")
    return formatted_date

def make_press_list(press_csv):
    pass ##########
def make_presentation_list(presentation_csv):
    
    format_list = ["Invited","Oral","Poster"]
    all_df = pd.read_csv(presentation_csv)
    doc.add_heading("Presentations (Japanese Titles were Translated to English)", level=1)
    for j,df in enumerate(format_list):
        df = all_df[all_df["Format"] == format_list[j]]
        df = df.sort_values(by = "Date", ascending = False)
        N = df.shape[0]
        doc.add_heading(format_list[j] + " Presentations"+f" ({N})", level=2)
        for i in range(N):
            data = df.iloc[i].astype(str)
            date,conference,venue,country_or_city,title,authors,format,notes = data[:8]
            formatted_date = format_date(date)
            p = doc.add_paragraph(f"{i+1}.  ")
            p = add_formatted_authors(p,authors)
            p.add_run(f' "{title}"').bold = True
            p.add_run(f" {conference}, {venue}, {country_or_city} ({formatted_date}).").add_break()
            if notes !="nan":
                comment = p.add_run(notes)
                print(notes)
                comment.font.bold = True
                # comment.font.underline = True
                comment.font.color.rgb = RGBColor.from_string("B10026")
                comment.add_break()
doc = Document("Publication_List_Template.docx")
# items = ["Heading", "Publications", "Presentations","Patents","Press","Funding","Awards"]
# for item in items:
#     doc.add_paragraph(item,style = "Heading 1")


    

make_publication_list("Publications.bib")
make_presentation_list("Presentations.csv")
make_funding_list("Funding.csv")
make_patent_list("Patents.csv")
make_award_list("Awards.csv")
# make_others_list("Others.csv")
# make_press_list("Press.csv")

doc.save("2_Publication_list.docx")

# %%
